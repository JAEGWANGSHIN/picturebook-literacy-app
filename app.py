from __future__ import annotations

# ── pptxgenjs 자동 설치 (Streamlit Cloud 대응) ──────────────────
import subprocess, sys, os

def _ensure_pptxgenjs():
    """make_pptx.js 와 같은 디렉터리에 pptxgenjs 를 로컬 설치."""
    app_dir = os.path.dirname(os.path.abspath(__file__))
    nm_path = os.path.join(app_dir, "node_modules", "pptxgenjs")
    if os.path.isdir(nm_path):
        return  # 이미 설치됨
    subprocess.run(
        ["npm", "install", "pptxgenjs"],
        cwd=app_dir,
        check=True,
        timeout=180,
        capture_output=True,
    )

try:
    _ensure_pptxgenjs()
except Exception:
    pass  # 설치 실패해도 앱은 실행 — PPT 기능만 비활성

import json, os, subprocess, tempfile
from datetime import datetime
from io import BytesIO

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt

st.set_page_config(
    page_title="AI 그림책 질문수업 설계기",
    page_icon="📚",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# ═══════════════════════════════════════════════════════════════════
# DB
# ═══════════════════════════════════════════════════════════════════
PICTUREBOOK_DB = [
    # ── 음운인식 ──────────────────────────────────────────────────
    {"id":"pb001","title":"말놀이 동시집","author":"최승호·방시혁","publisher":"비룡소",
     "theme":["음운인식","어휘"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"동물 이름과 의성어로 구성된 말놀이 동시 모음",
     "literacy_elements":["음운인식","어휘"],
     "reason":"한국어 음소·음절 인식, 운율 체험"},
    {"id":"pb002","title":"수수께끼야 놀자","author":"이상교","publisher":"보림",
     "theme":["음운인식","어휘"],"grade":["유치원","초등 1학년"],
     "summary":"수수께끼 형식으로 의미와 소리를 연결",
     "literacy_elements":["음운인식","어휘"],
     "reason":"소리 단서로 단어 맞추기; 파닉스 연결"},
    {"id":"pb003","title":"Brown Bear, Brown Bear","author":"Bill Martin Jr.","publisher":"Henry Holt",
     "theme":["음운인식","어휘"],"grade":["유치원","초등 1학년"],
     "summary":"색깔과 동물 이름이 반복되는 패턴 그림책",
     "literacy_elements":["음운인식","어휘"],
     "reason":"반복 패턴으로 운율·예측 읽기"},
    # ── 어휘 ──────────────────────────────────────────────────────
    {"id":"pb004","title":"단어 수집가","author":"Peter H. Reynolds","publisher":"Candlewick",
     "theme":["어휘","정체성"],"grade":["초등 1학년","초등 2학년","초등 3학년"],
     "summary":"소년이 세상의 단어들을 수집하는 이야기",
     "literacy_elements":["어휘","음운인식","정체성"],
     "reason":"단어 가치 인식; 어휘 수집 동기 부여"},
    {"id":"pb005","title":"알사탕","author":"백희나","publisher":"책읽는곰",
     "theme":["어휘","감정 이해","가족"],"grade":["초등 1학년","초등 2학년"],
     "summary":"신비한 사탕을 먹으면 주변의 소리가 들린다",
     "literacy_elements":["어휘","감정이해"],
     "reason":"감각어·감정 어휘 풍부; 서정적 표현"},
    {"id":"pb006","title":"구름빵","author":"백희나","publisher":"한솔수북",
     "theme":["어휘","가족","상상력"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"비 오는 날 구름으로 빵을 만들어 날아다니는 이야기",
     "literacy_elements":["어휘","이야기이해","상상력"],
     "reason":"요리·자연 어휘; 판타지 어휘 확장"},
    {"id":"pb007","title":"수박 수영장","author":"안녕달","publisher":"창비",
     "theme":["어휘","배경지식","다양성 존중","상상력"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"커다란 수박 속 수영장에서 동네 사람들이 함께 노는 상상",
     "literacy_elements":["어휘","배경지식","상상력"],
     "reason":"여름·공동체·감각 어휘; 배경지식 확장"},
    {"id":"pb027","title":"The Very Hungry Caterpillar","author":"Eric Carle","publisher":"Puffin Books",
     "theme":["배경지식","어휘"],"grade":["유치원","초등 1학년"],
     "summary":"배고픈 애벌레가 다양한 음식을 먹으며 성장하는 이야기",
     "literacy_elements":["이야기 재구성","배경지식","어휘"],
     "reason":"요일·음식·변태 순서 재구성; 반복 구조"},
    # ── 배려 ──────────────────────────────────────────────────────
    {"id":"pb008","title":"이상한 손님","author":"백희나","publisher":"책읽는곰",
     "theme":["어휘","추론하기","배려"],"grade":["초등 1학년","초등 2학년"],
     "summary":"비 오는 날 하늘 나라에서 길 잃은 아이가 찾아온 이야기",
     "literacy_elements":["어휘","추론하기","배려"],
     "reason":"감정 어휘·비유 표현; 상황 맥락 추론; 배려"},
    {"id":"pb015","title":"무지개 물고기","author":"마르쿠스 피스터","publisher":"시공주니어",
     "theme":["어휘","배려","친구 관계"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"아름다운 비늘을 나눠주며 친구를 사귀는 물고기 이야기",
     "literacy_elements":["어휘","감정이해","배려","친구 관계"],
     "reason":"바닷속 어휘; 나눔과 배려; 친구 관계 탐색"},
    {"id":"pb016","title":"강아지똥","author":"권정생","publisher":"길벗어린이",
     "theme":["자존감","감정 이해","배려"],"grade":["초등 1학년","초등 2학년","초등 3학년"],
     "summary":"아무도 거들떠보지 않던 강아지똥이 민들레의 거름이 된다",
     "literacy_elements":["어휘","감정이해","배려","자존감"],
     "reason":"자연 어휘·존재 가치; 배려와 자존감 통합"},
    {"id":"pb026","title":"으뜸 헤엄이(Swimmy)","author":"Leo Lionni","publisher":"시공주니어",
     "theme":["친구 관계","용기","배려"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"혼자 헤엄치는 물고기가 친구들과 힘을 합치는 이야기",
     "literacy_elements":["이야기 재구성","감정이해","용기","배려","친구 관계"],
     "reason":"협력과 배려; 용기; 친구 관계; 이야기 재구성"},
    # ── 자존감 ────────────────────────────────────────────────────
    {"id":"pb009","title":"내 귀는 짝짝이","author":"율리 슈타르크","publisher":"비룡소",
     "theme":["자존감","다양성 존중","정체성"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"귀가 다른 토끼가 자신을 있는 그대로 받아들이는 이야기",
     "literacy_elements":["감정이해","자존감","정체성"],
     "reason":"신체 다양성 수용; 자기 긍정 감정 어휘"},
    {"id":"pb010","title":"너는 특별하단다","author":"맥스 루케이도","publisher":"고슬링",
     "theme":["자존감","정체성"],"grade":["유치원","초등 1학년","초등 2학년","초등 3학년"],
     "summary":"작은 나무 사람 펀치넬로가 자신의 가치를 깨닫는 이야기",
     "literacy_elements":["감정이해","이야기이해","자존감","정체성"],
     "reason":"자존감 언어; 평가와 자기 가치 탐색"},
    {"id":"pb013","title":"100만 번 산 고양이","author":"사노 요코","publisher":"비룡소",
     "theme":["자존감","정체성","감정 이해"],"grade":["초등 2학년","초등 3학년","초등 4학년"],
     "summary":"100만 번을 살면서 진정한 사랑을 깨달은 고양이 이야기",
     "literacy_elements":["이야기이해","감정이해","자존감","정체성"],
     "reason":"삶과 사랑의 의미; 감정 변화 추적"},
    {"id":"pb018","title":"꽃들에게 희망을","author":"트리나 폴러스","publisher":"분도출판사",
     "theme":["자존감","정체성","용기"],"grade":["초등 2학년","초등 3학년","초등 4학년"],
     "summary":"애벌레가 자아를 찾아 나비가 되는 우화",
     "literacy_elements":["어휘","추론하기","자존감","정체성","용기"],
     "reason":"삶의 의미 어휘; 변화·희망; 자존감·용기"},
    {"id":"pb020","title":"괜찮아","author":"최숙희","publisher":"웅진주니어",
     "theme":["자존감","다양성 존중","감정 이해"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"서로 다른 모습도 괜찮다는 자기 수용 이야기",
     "literacy_elements":["감정이해","어휘","자존감"],
     "reason":"자기 수용·다양성 감정 표현 어휘"},
    {"id":"pb022","title":"나쁜 어린이 표","author":"황선미","publisher":"웅진주니어",
     "theme":["자존감","감정 이해"],"grade":["초등 1학년","초등 2학년","초등 3학년"],
     "summary":"잘못을 저지른 어린이가 표를 붙이고 다니는 이야기",
     "literacy_elements":["이야기이해","감정이해","자존감"],
     "reason":"원인-결과; 인물 내면 변화; 자존감 회복"},
    # ── 정체성 ────────────────────────────────────────────────────
    {"id":"pb011","title":"중요한 사실","author":"마가렛 와이즈 브라운","publisher":"시공주니어",
     "theme":["정체성","어휘"],"grade":["초등 1학년","초등 2학년"],
     "summary":"사물의 가장 중요한 특성에 대해 이야기하는 철학적 그림책",
     "literacy_elements":["어휘","추론하기","정체성"],
     "reason":"핵심 특성 파악; 자아에 대한 질문 생성"},
    {"id":"pb033","title":"선생님이 나를 모르면","author":"이상교","publisher":"보림",
     "theme":["정체성","의사소통"],"grade":["초등 1학년"],
     "summary":"아이가 선생님에게 자신을 소개하는 이야기",
     "literacy_elements":["이야기이해","정체성","의사소통"],
     "reason":"나에 대한 질문 생성; 자기 이해 촉진"},
    {"id":"pb034","title":"나는 어떻게 생겨났을까?","author":"과학그림책","publisher":"비룡소",
     "theme":["배경지식","정체성"],"grade":["초등 1학년","초등 2학년"],
     "summary":"탄생의 과학적 사실을 어린이 눈높이로 설명",
     "literacy_elements":["배경지식","정체성"],
     "reason":"배경지식 궁금증에서 질문 생성 자연 유도"},
    # ── 친구 관계 ─────────────────────────────────────────────────
    {"id":"pb012","title":"고슴도치 X","author":"에밀리 그레이벳","publisher":"웅진주니어",
     "theme":["친구 관계","감정 이해","의사소통"],"grade":["초등 1학년","초등 2학년"],
     "summary":"편지를 쓰다가 계속 틀려서 X로 지워나가는 고슴도치 이야기",
     "literacy_elements":["감정이해","이야기이해","친구 관계","의사소통"],
     "reason":"감정 표현의 어려움; 친구에게 마음 전하기; 의사소통"},
    {"id":"pb032","title":"Voices in the Park","author":"앤서니 브라운","publisher":"DK Publishing",
     "theme":["다양성 존중","친구 관계","의사소통"],"grade":["초등 2학년","초등 3학년","초등 4학년"],
     "summary":"같은 공원 방문을 4명의 서로 다른 목소리로 이야기",
     "literacy_elements":["추론하기","이야기이해","친구 관계","의사소통"],
     "reason":"관점 추론; 같은 사건의 다른 해석; 친구 관계"},
    # ── 의사소통 ──────────────────────────────────────────────────
    {"id":"pb017","title":"돼지책","author":"앤서니 브라운","publisher":"웅진주니어",
     "theme":["가족","다양성 존중","의사소통"],"grade":["초등 2학년","초등 3학년","초등 4학년"],
     "summary":"혼자 집안일을 하던 엄마가 집을 나가고 가족이 돼지로 변한다",
     "literacy_elements":["이야기이해","추론하기","의사소통"],
     "reason":"인물 동기·감정 변화 추적; 시각 상징 분석; 가족 의사소통"},
    {"id":"pb019","title":"고구마구마","author":"사이다","publisher":"반달",
     "theme":["어휘","의사소통","감정 이해"],"grade":["유치원","초등 1학년"],
     "summary":"고구마가 '구마'라고만 말하는 반복 언어유희 그림책",
     "literacy_elements":["어휘","음운인식","의사소통"],
     "reason":"파닉스 연결; 반복 패턴; 말로 마음 전하기"},
    {"id":"pb024","title":"지각대장 존","author":"John Burningham","publisher":"비룡소",
     "theme":["상상력","의사소통"],"grade":["초등 1학년","초등 2학년"],
     "summary":"매일 지각하는 존의 기상천외한 이유",
     "literacy_elements":["이야기이해","추론하기","상상력","의사소통"],
     "reason":"사실과 상상 구별; 인물 관점 이해; 의사소통"},
    # ── 용기 ──────────────────────────────────────────────────────
    {"id":"pb014","title":"빈집에 온 손님","author":"김유경","publisher":"웅진주니어",
     "theme":["두려움","용기","상상력"],"grade":["초등 1학년","초등 2학년"],
     "summary":"홀로 집을 지키던 아이가 상상 속 손님을 맞이하는 이야기",
     "literacy_elements":["추론하기","감정이해","용기","상상력"],
     "reason":"두려움 극복과 용기; 감정 탐색"},
    {"id":"pb031","title":"Two Bad Ants","author":"Chris Van Allsburg","publisher":"Houghton Mifflin",
     "theme":["상상력","용기"],"grade":["초등 2학년","초등 3학년"],
     "summary":"두 개미가 설탕 그릇으로 모험을 떠나는 이야기",
     "literacy_elements":["추론하기","상상력","용기"],
     "reason":"개미 시점 시각 추론의 정수; 용기와 모험"},
    # ── 추론하기 ──────────────────────────────────────────────────
    {"id":"pb023","title":"선물","author":"이수지","publisher":"비룡소",
     "theme":["상상력","감정 이해","추론하기"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"눈이 오는 날의 감동을 글 없이 그림만으로 표현",
     "literacy_elements":["추론하기","감정이해","상상력"],
     "reason":"무자(wordless) 그림책; 그림만으로 추론하는 최적 텍스트"},
    {"id":"pb028","title":"Where the Wild Things Are","author":"Maurice Sendak","publisher":"HarperCollins",
     "theme":["감정 조절","상상력","추론하기"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"맥스가 상상의 세계로 여행하고 집으로 돌아오는 이야기",
     "literacy_elements":["이야기 재구성","감정이해","추론하기","상상력"],
     "reason":"여행 구조 재구성; 상징 추론 전형"},
    {"id":"pb021","title":"7년 동안의 잠","author":"박완서","publisher":"웅진주니어",
     "theme":["배경지식","감정 이해","추론하기"],"grade":["초등 2학년","초등 3학년"],
     "summary":"흉년 든 개미마을에 나타난 번데기를 둘러싼 이야기",
     "literacy_elements":["이야기이해","배경지식","감정이해","추론하기"],
     "reason":"생태 배경지식; 발단·전개·결말 구조 분석"},
    # ── 상상력 ────────────────────────────────────────────────────
    {"id":"pb025","title":"100층짜리 집","author":"이와이 도시오","publisher":"북뱅크",
     "theme":["배경지식","상상력"],"grade":["유치원","초등 1학년"],
     "summary":"주인공이 100층까지 올라가며 여러 동물을 만나는 이야기",
     "literacy_elements":["이야기이해","배경지식","상상력"],
     "reason":"순서·수 개념; 동물 생태 배경지식; 상상력"},
    {"id":"pb030","title":"왜냐하면(Because)","author":"Mo Willems","publisher":"Hyperion",
     "theme":["상상력","배경지식"],"grade":["초등 1학년","초등 2학년"],
     "summary":"연쇄적 원인-결과로 이어지는 이야기",
     "literacy_elements":["이야기이해","추론하기","상상력"],
     "reason":"'왜?' 질문 구조를 시각적으로 보여줌; 상상력"},
    # ── 기타 ──────────────────────────────────────────────────────
    {"id":"pb029","title":"In My Heart","author":"Jo Witek","publisher":"Abrams Appleseed",
     "theme":["감정 이해","감정 조절"],"grade":["유치원","초등 1학년"],
     "summary":"다양한 감정을 신체 감각으로 묘사하는 그림책",
     "literacy_elements":["감정이해","어휘"],
     "reason":"감정 어휘 10가지 명시적 학습"},
    {"id":"pb035","title":"The Invisible String","author":"Patrice Karst","publisher":"DeVorss",
     "theme":["감정 이해","두려움","가족"],"grade":["유치원","초등 1학년","초등 2학년"],
     "summary":"사랑하는 사람과의 보이지 않는 연결 이야기",
     "literacy_elements":["감정이해"],
     "reason":"분리불안·연결감 감정; 저학년 적합"},
]


def db_search(theme="", grade=""):
    return [b for b in PICTUREBOOK_DB if
            ((not theme) or any(theme in t for t in b["theme"])) and
            ((not grade) or grade in b["grade"])]

def db_get_by_title(title):
    return next((b for b in PICTUREBOOK_DB if b["title"] == title), None)

def db_all_themes():
    t = set()
    for b in PICTUREBOOK_DB: t.update(b["theme"])
    return sorted(t)

# ═══════════════════════════════════════════════════════════════════
# CSS
# ═══════════════════════════════════════════════════════════════════
CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;500;700&display=swap');

/* ── 전체 리셋 ── */
html, body,
[data-testid="stAppViewContainer"],
[data-testid="stAppViewContainer"] > .main {
  background-color: #F7F8FA !important;
  font-family: 'Noto Sans KR', sans-serif !important;
}
[data-testid="stSidebar"] { display: none !important; }
.main .block-container {
  max-width: 800px !important;
  padding: 0 1.2rem 3rem !important;
}
* { font-family: 'Noto Sans KR', sans-serif !important; }

/* ── 상단 헤더바 ── */
.top-bar {
  background: #1A1A2E;
  color: white;
  padding: .9rem 1.6rem;
  margin: 0 -1.2rem 1.6rem;
  display: flex;
  align-items: center;
  gap: 10px;
}
.top-bar-icon { font-size: 1.3rem; }
.top-bar-title {
  font-size: 1.05rem;
  font-weight: 700;
  color: white;
  letter-spacing: -.3px;
}
.top-bar-sub {
  font-size: .78rem;
  color: #A0A8C0;
  margin-left: auto;
}


.section-title {
  font-size: .78rem;
  font-weight: 700;
  color: #6B7280;
  text-transform: uppercase;
  letter-spacing: .8px;
  margin: 1.2rem 0 .7rem;
  padding: .55rem .9rem;
  display: flex;
  align-items: center;
  gap: 8px;
  border-radius: 8px;
}
.section-title::before {
  content: '';
  display: inline-block;
  width: 4px; height: 16px;
  background: #1A1A2E;
  border-radius: 3px;
  flex-shrink: 0;
}
/* 파스텔 배경을 section-title 자체에 적용 */
.section-title.st-blue   { background: #EFF6FF; color: #1D4ED8; }
.section-title.st-purple { background: #F5F3FF; color: #6D28D9; }
.section-title.st-green  { background: #F0FDF4; color: #065F46; }



/* ── 구분선 ── */
.divider {
  height: 1px;
  background: #E8EAF0;
  border: none;
  margin: 1rem 0;
}

/* 이전 호환용 */
.sec-label {
  font-size: .72rem !important;
  font-weight: 700 !important;
  color: #6B7280 !important;
  text-transform: uppercase;
  letter-spacing: .8px;
  margin: 0 0 .7rem !important;
  display: flex !important;
  align-items: center;
  gap: 6px;
}

/* ── 폼 ── */
[data-testid="stSelectbox"] > div > div,
[data-testid="stTextInput"] > div > div > input,
[data-testid="stTextArea"] > div > div > textarea {
  border-radius: 8px !important;
  border: 1px solid #D1D5DB !important;
  background: white !important;
  font-size: .88rem !important;
}
/* 라벨 caption과 위젯 사이 간격 최소화 */
[data-testid="stCaptionContainer"] {
  margin-bottom: -10px !important;
  padding-bottom: 0 !important;
}
[data-testid="stSelectbox"] > div > div:focus-within,
[data-testid="stTextInput"] > div > div > input:focus,
[data-testid="stTextArea"] > div > div > textarea:focus {
  border-color: #1A1A2E !important;
  box-shadow: 0 0 0 2px #1A1A2E18 !important;
}

/* ── 책 카드 ── */
.book-card {
  background: #F9FAFB;
  border: 1px solid #E5E7EB;
  border-radius: 10px;
  padding: .8rem 1rem;
  margin: .4rem 0;
  display: flex;
  gap: 10px;
  align-items: flex-start;
}
.book-card.selected {
  background: #EEF2FF;
  border-color: #6366F1;
}
.bc-icon  { font-size: 1.3rem; flex-shrink: 0; }
.bc-title { font-weight: 700; color: #111827; font-size: .88rem; }
.bc-meta  { color: #6B7280; font-size: .75rem; margin-top: 2px; line-height: 1.45; }
.bc-tags  { margin-top: 5px; display: flex; flex-wrap: wrap; gap: 4px; }
.tag {
  background: #F3F4F6;
  color: #374151;
  border: 1px solid #E5E7EB;
  border-radius: 4px;
  padding: 1px 7px;
  font-size: .67rem;
  font-weight: 500;
}

/* ── AI 추천 버튼 ── */
/* 추천 책 버튼: secondary 스타일 그대로 활용, 선택 시 ✅ 접두어로 구분 */

/* ── 질문 카드 ── */
.q-grid {
  display: grid;
  grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
  gap: 8px;
  margin-top: .6rem;
}
.q-card {
  background: white;
  border: 1px solid #E5E7EB;
  border-radius: 8px;
  padding: .75rem .85rem;
  transition: box-shadow .15s;
}
.q-card:hover { box-shadow: 0 4px 12px #0000000D; }
.q-card .qt {
  font-size: .65rem;
  font-weight: 700;
  border-radius: 4px;
  padding: 2px 7px;
  display: inline-block;
  margin-bottom: 5px;
}
.q-card .qt.사실  { background: #EFF6FF; color: #1D4ED8; }
.q-card .qt.추론  { background: #F5F3FF; color: #6D28D9; }
.q-card .qt.평가  { background: #F0FDF4; color: #166534; }
.q-card .qt.감정  { background: #FFF1F2; color: #9F1239; }
.q-card .qt.작가  { background: #FFF7ED; color: #C2410C; }
.q-card .qt.삶연결{ background: #ECFDF5; color: #065F46; }
.q-card .qtext { font-size: .82rem; color: #374151; line-height: 1.5; }

/* ── 지도안 테이블 ── */
.lp-table { width:100%; border-collapse:collapse; margin-bottom:14px; font-size:.82rem; }
.lp-table th {
  padding:8px 10px; font-size:.72rem; font-weight:700;
  text-align:center; background:#F9FAFB; color:#374151; border:1px solid #E5E7EB;
}
.lp-table td { padding:9px 11px; border:1px solid #E5E7EB; vertical-align:top; line-height:1.6; }
.lp-stage {
  font-size:.72rem; font-weight:700; text-align:center;
  border-radius:6px; padding:3px 8px; display:inline-block; white-space:nowrap;
}
.lp-dur { font-size:.72rem; color:#6B7280; text-align:center; }
.lp-teacher { color:#1D4ED8; }
.lp-student { color:#166534; }
.lp-note   { color:#92400E; background:#FFFBEB; border-radius:6px; padding:4px 7px; font-size:.78rem; }
/* ── 활동 카드 ── */
.act-card {
  background: white;
  border: 1px solid #E5E7EB;
  border-radius: 12px;
  padding: 1.1rem 1.2rem;
  margin-bottom: 10px;
}
.act-header {
  display: flex;
  align-items: center;
  gap: 10px;
  margin-bottom: .7rem;
  padding-bottom: .6rem;
  border-bottom: 1px solid #F3F4F6;
}
.act-step-badge {
  background: #1A1A2E;
  color: white;
  border-radius: 6px;
  padding: 3px 10px;
  font-size: .72rem;
  font-weight: 700;
}
.act-title { font-size: 1rem; font-weight: 700; color: #111827; }
.act-duration {
  margin-left: auto;
  background: #F0FDF4;
  color: #166534;
  border: 1px solid #BBF7D0;
  border-radius: 20px;
  padding: 2px 10px;
  font-size: .72rem;
  font-weight: 600;
}
.act-row { display: flex; gap: 6px; margin-bottom: 5px; align-items: flex-start; }
.act-label {
  font-size: .7rem; font-weight: 700; color: #6B7280;
  min-width: 52px; padding-top: 1px;
}
.act-val { font-size: .83rem; color: #374151; line-height: 1.5; }
.act-question {
  background: #EFF6FF;
  border-radius: 8px;
  padding: .55rem .8rem;
  margin-top: .6rem;
  font-size: .83rem;
  color: #1D4ED8;
  font-style: italic;
}
/* ── 평가 카드 ── */
.eval-table {
  width: 100%;
  border-collapse: collapse;
  margin-bottom: 12px;
  font-size: .82rem;
}
.eval-table th {
  padding: 7px 10px;
  font-weight: 700;
  font-size: .72rem;
  text-align: left;
}
.eval-table td { padding: 8px 10px; border-bottom: 1px solid #F3F4F6; vertical-align: top; }
.eval-table tr:last-child td { border-bottom: none; }
.th-area { background: #F9FAFB; color: #374151; width: 20%; }
.th-good { background: #F0FDF4; color: #166534; width: 27%; }
.th-ok   { background: #FFFBEB; color: #92400E; width: 27%; }
.th-need { background: #EFF6FF; color: #1D4ED8; width: 26%; }
.td-area { background: #F9FAFB; font-weight: 600; color: #374151; }
/* ── 결과 섹션 (커스텀 — expander 미사용) ── */
.result-section {
  background: white;
  border: 1px solid #E5E7EB;
  border-radius: 10px;
  margin-bottom: 8px;
  overflow: hidden;
}
.result-section-body {
  padding: .2rem 1rem .8rem;
  border-top: 1px solid #F3F4F6;
}

/* ── 전체 생성 버튼 (특별 강조) ── */
[data-testid="baseButton-primary"].generate-all-btn,
.generate-all [data-testid="baseButton-primary"] {
  background: linear-gradient(135deg, #1A1A2E, #3730A3) !important;
  font-size: 1rem !important;
  padding: .8rem 0 !important;
  letter-spacing: .3px;
}
/* ── 메인 버튼 ── */
[data-testid="baseButton-primary"] {
  background: #1A1A2E !important;
  border: none !important;
  border-radius: 8px !important;
  font-size: .9rem !important;
  font-weight: 600 !important;
  color: white !important;
  transition: opacity .15s !important;
  letter-spacing: -.2px;
}
[data-testid="baseButton-primary"]:hover { opacity: .88 !important; }
[data-testid="baseButton-secondary"] {
  border-radius: 8px !important;
  border: 1px solid #D1D5DB !important;
  background: white !important;
  font-size: .88rem !important;
  font-weight: 500 !important;
  color: #374151 !important;
  transition: border-color .15s, background .15s !important;
}
[data-testid="baseButton-secondary"]:hover {
  border-color: #6B7280 !important;
  background: #F9FAFB !important;
}

/* ── 탭 ── */
[data-testid="stTabs"] [role="tab"] {
  font-size: .88rem !important;
  font-weight: 500 !important;
  color: #6B7280 !important;
  padding: .5rem .9rem !important;
}
[data-testid="stTabs"] [role="tab"][aria-selected="true"] {
  color: #111827 !important;
  font-weight: 700 !important;
  border-bottom: 2px solid #1A1A2E !important;
}

/* ── PPT 버튼 ── */
.ppt-wrap [data-testid="baseButton-secondary"] {
  background: #F0FDF4 !important;
  border: 1px solid #86EFAC !important;
  color: #166534 !important;
}

/* ── 반응형 ── */
@media(max-width:600px){
  .main .block-container { padding: 0 .6rem 2rem !important; }
  .top-bar { margin: 0 -.6rem 1.2rem; padding: .8rem 1rem; }
  .q-grid { grid-template-columns: 1fr; }
}
</style>
"""
PPTX_SCRIPT = os.path.join(os.path.dirname(__file__), "make_pptx.js")

# ═══════════════════════════════════════════════════════════════════
# OpenAI 헬퍼
# ═══════════════════════════════════════════════════════════════════
def get_client():
    key = st.secrets.get("OPENAI_API_KEY", None) or os.getenv("OPENAI_API_KEY")
    if not key:
        st.error("OpenAI API 키가 없습니다. Streamlit Secrets에 OPENAI_API_KEY를 등록하세요.")
        st.stop()
    return OpenAI(api_key=key)

def chat(system: str, user: str, max_tokens: int = 1200) -> str:
    client = get_client()
    r = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"system","content":system},
                  {"role":"user","content":user}],
        temperature=0.7,
        max_tokens=max_tokens,
    )
    return r.choices[0].message.content or ""

# ── AI 추천 ──────────────────────────────────────────────────────
def ai_recommend_books(situation: str) -> list[dict]:
    db_titles = "\n".join(
        f"- {b['title']} ({b['author']}): {b['summary']} [주제: {', '.join(b['theme'])}]"
        for b in PICTUREBOOK_DB
    )
    resp = chat(
        "당신은 초등 그림책 전문가입니다. 교사의 상황에 맞는 그림책을 추천하세요.",
        f"""교사 상황: {situation}

아래 그림책 목록 중 가장 적합한 책 4권을 추천하세요.
반드시 목록에 있는 제목만 사용하고, JSON 배열로만 응답하세요.
형식: [{{"title":"제목","reason":"한 문장 추천 이유"}}]

그림책 목록:
{db_titles}""",
        max_tokens=600,
    )
    try:
        import re
        m = re.search(r'\[.*\]', resp, re.DOTALL)
        return json.loads(m.group()) if m else []
    except Exception:
        return []

# ── 질문 생성 ─────────────────────────────────────────────────────
def gen_questions(grade, theme, book, book_info, book_ctx_extra: str = "") -> dict:
    if book_info:
        pub_info = f" ({book_info.get('publisher','')})" if book_info.get('publisher') else ""
        book_ctx = (
            f"작가: {book_info['author']}{pub_info}\n"
            f"줄거리: {book_info['summary']}\n"
            f"이 책의 문해력 요소: {', '.join(book_info.get('literacy_elements', []))}"
        )
    elif book_ctx_extra:
        book_ctx = f"[교사 제공 책 정보]\n{book_ctx_extra}"
    else:
        book_ctx = ""

    theme_instruction = (
        f"수업 주제는 '{theme}'입니다. "
        f"질문의 상당수(최소 3~4개)는 반드시 '{theme}' 관련 내용을 포함해야 합니다. "
        f"예를 들어 주제가 '음운인식'이라면 소리·음절·운율·글자 등에 대한 질문을 꼭 넣어 주세요. "
        f"나머지 질문은 그림책 내용(이야기·감정·인물 등)과 연결합니다."
    )

    resp = chat(
        "당신은 초등 그림책 질문 수업 전문가입니다. 반드시 JSON으로만 응답하세요.",
        f"""그림책: {book}
{book_ctx}
학년: {grade}
{theme_instruction}

읽기 전 질문 3개, 읽는 중 질문 5개, 읽은 후 질문 5개를 만들어 주세요.
각 질문은 사실/추론/평가/감정/작가/삶연결 중 하나의 유형을 가집니다.

반드시 아래 JSON 형식으로만 응답하세요:
{{"before":[{{"type":"유형","text":"질문"}},...],
  "during":[{{"type":"유형","text":"질문"}},...],
  "after" :[{{"type":"유형","text":"질문"}},...] }}""",
        max_tokens=1200,
    )
    try:
        import re
        m = re.search(r'\{.*\}', resp, re.DOTALL)
        return json.loads(m.group()) if m else {}
    except Exception:
        return {}

# ── 활동 생성 ─────────────────────────────────────────────────────
def gen_activities(grade, theme, book, lesson_time, student_context, book_ctx_extra: str = "") -> str:
    extra = f"\n[교사 제공 책 정보]\n{book_ctx_extra}" if book_ctx_extra else ""
    return chat(
        "초등 수업 설계 전문가입니다. 반드시 JSON으로만 응답하세요.",
        f"""학년:{grade} / 주제:{theme} / 그림책:{book} / 시간:{lesson_time}
학생특성:{student_context or '없음'}{extra}

3가지 활동(도입/중심/정리)을 설계하세요. 총 수업 시간: {lesson_time}

반드시 아래 JSON 형식으로만 응답하세요:
[
  {{
    "step": "도입",
    "title": "활동명",
    "duration": "10분",
    "goal": "활동 목표 한 문장",
    "materials": "준비물 (쉼표 구분)",
    "process": ["진행 순서 1", "진행 순서 2", "진행 순서 3"],
    "question": "교사 핵심 발문"
  }},
  {{ "step": "중심", ... }},
  {{ "step": "정리", ... }}
]""",
        max_tokens=1400,
    )

# ── 지도안 생성 ───────────────────────────────────────────────────
def gen_lessonplan(grade, theme, book, lesson_time, student_context,
                   book_ctx_extra: str = "", questions: dict = None) -> str:
    extra = f"\n[교사 제공 책 정보]\n{book_ctx_extra}" if book_ctx_extra else ""

    # 질문 카드 요약 — 지도안에 반영
    q_ctx = ""
    if questions:
        before = questions.get("before", [])
        during = questions.get("during", [])
        after  = questions.get("after",  [])
        def fmt(qs): return " / ".join(q.get("text","") if isinstance(q,dict) else str(q) for q in qs[:3])
        q_ctx = f"""
[생성된 질문 카드 — 지도안에 반드시 반영하세요]
읽기 전 질문: {fmt(before)}
읽는 중 질문: {fmt(during)}
읽은 후 질문: {fmt(after)}
"""

    return chat(
        "초등 수업 설계 전문가입니다. 반드시 JSON으로만 응답하세요.",
        f"""학년:{grade} / 주제:{theme} / 그림책:{book} / 시간:{lesson_time}
학생특성:{student_context or '없음'}{extra}{q_ctx}

[시간 배분 — 고정]
- 도입: 5분
- 전개: {lesson_time.replace("분","").strip() if "분" in lesson_time else "30"}분에서 10을 뺀 분 (활동1·활동2, 필요 시 활동3)
- 정리: 5분

[작성 원칙 — 반드시 지키세요]
1. teacher(교사 발문/활동): 질문은 번호를 붙여 각각 별도 문자열로, "questions" 배열에 저장
2. student(학생 활동): 각 교사 질문에 대한 구체적 학생 반응/예상 답변을 포함. 단순 행동만 쓰지 말고
   "A1: 예상 답변 예시 / A2: 예상 답변 예시" 형태로 상세히 작성
3. note(유의점): 이 그림책·주제에 특화된 유의점. "학생을 격려한다" 같은 뻔한 표현 금지.
   예) "「{book}」의 경우 ~한 장면에서 아이들이 ~반응을 보일 수 있으므로 ~" 처럼 구체적으로

JSON 형식:
[
  {{
    "stage": "도입",
    "duration": "5분",
    "rows": [
      {{
        "questions": ["읽기 전 질문1", "읽기 전 질문2"],
        "teacher": "교사 도입 활동 설명 (질문 제시 방법 포함)",
        "student": "A1: 예시 답변 / A2: 예시 답변. 분위기·참여 방식 설명",
        "note": "이 그림책 특화 유의점"
      }}
    ]
  }},
  {{
    "stage": "전개",
    "duration": "X분",
    "rows": [
      {{
        "activity": "활동1: 대화형 읽기",
        "questions": ["읽는 중 질문1", "읽는 중 질문2", "읽는 중 질문3"],
        "teacher": "그림책 읽어주며 질문 제시 방법",
        "student": "A1: 구체적 예상 답변 / A2: 구체적 예상 답변 / 학생 상호작용 방식",
        "note": "이 그림책 특화 유의점"
      }},
      {{
        "activity": "활동2: 독후 표현",
        "questions": [],
        "teacher": "독후 표현 활동 안내",
        "student": "구체적인 활동 내용, 표현 방식, 예상 결과물",
        "note": "이 그림책 특화 유의점"
      }}
    ]
  }},
  {{
    "stage": "정리",
    "duration": "5분",
    "rows": [
      {{
        "questions": ["읽은 후 질문1", "읽은 후 질문2"],
        "teacher": "정리 활동 및 질문 제시",
        "student": "A1: 구체적 예상 답변 / A2: 구체적 예상 답변. 성찰 방식",
        "note": "이 그림책 특화 유의점"
      }}
    ]
  }}
]""",
        max_tokens=2000,
    )

# ── 평가+안내문 생성 ──────────────────────────────────────────────
def gen_eval_parent(grade, theme, book, book_ctx_extra: str = "") -> str:
    extra = f"\n[교사 제공 책 정보]\n{book_ctx_extra}" if book_ctx_extra else ""
    return chat(
        "초등 수업 평가 전문가입니다. 반드시 JSON으로만 응답하세요.",
        f"""학년:{grade} / 주제:{theme} / 그림책:{book}{extra}

아래 JSON 형식으로 정확히 응답하세요.
평가 서술 원칙: 모든 단계를 긍정적·성장 지향적으로 씁니다.
"부족하다", "못한다", "어렵다" 같은 부정 표현 금지.
노력 필요는 반드시 "교사·친구의 도움을 받아 ～할 수 있다" 형태로.
반드시 "~합니다", "~있습니다" 같은 합쇼체 금지. 모든 서술어는 "~이다", "~한다", "~할 수 있다" 형태로 작성.

{{
  "criteria": [
    {{
      "area": "평가 영역명",
      "good": "잘함 서술 (스스로 능숙하게)",
      "ok": "보통 서술 (대체로 할 수 있음)",
      "needs": "노력 필요 서술 (도움을 받아 ~할 수 있다)"
    }}
  ],
  "self_eval": ["자기평가 문항1", "자기평가 문항2", "자기평가 문항3"],
  "parent_letter": "학부모 안내문 전체 텍스트 (그림책·주제 소개 + 가정 대화 질문 3개 + 따뜻한 어조)"
}}

criteria는 4개, self_eval은 3개.""",
        max_tokens=900,
    )



# ── 학생 활동지 생성 ──────────────────────────────────────────────
def gen_worksheet(grade, theme, book, book_info, questions: dict = None,
                  book_ctx_extra: str = "") -> str:
    """AI가 HTML 활동지를 생성"""
    extra = f"\n[교사 제공 책 정보]\n{book_ctx_extra}" if book_ctx_extra else ""

    # 질문 카드에서 핵심 질문 추출
    q_sample = ""
    if questions:
        all_qs = []
        for section in ["before","during","after"]:
            for q in questions.get(section, []):
                all_qs.append(q.get("text","") if isinstance(q,dict) else str(q))
        if all_qs:
            q_sample = "\n".join(f"- {q}" for q in all_qs[:5])

    pub = book_info.get("publisher","") if book_info else ""
    author = book_info.get("author","") if book_info else ""
    book_info_str = f"{author} / {pub}" if pub else author

    resp = chat(
        "초등 수업 활동지 디자이너입니다. 반드시 완전한 HTML만 출력하세요. 다른 텍스트 없이.",
        f"""학년:{grade} / 주제:{theme} / 그림책:{book} ({book_info_str})
학생 활동지를 만들어 주세요.{extra}

참고 질문 (일부 활용 가능):
{q_sample}

[디자인 원칙]
- A4 1장 (210×297mm), 인쇄 가능한 HTML
- 구글 폰트 Nanum Gothic + Gaegu 사용
- 파스텔 컬러, 이모지 적극 활용
- 학생이 흥미를 느낄 수 있는 세련되고 귀여운 디자인
- 그림책 제목과 주제가 상단에 크게 표시
- 활동 섹션: 읽기 전(1문항) + 읽으면서(2문항) + 읽은 후(2문항) + 나만의 그림/표현 공간 1개
- 각 문항 아래 충분한 답변 공간 (점선 또는 빈 박스)
- 하단에 이름·날짜 기입란

아래 조건을 반드시 지키세요:
1. <html>부터 </html>까지 완전한 HTML 출력
2. 외부 폰트는 Google Fonts CDN만 사용
3. print용 CSS 포함 (@media print)
4. 배경색은 인쇄 시 보이도록 -webkit-print-color-adjust: exact 설정
5. 전체 높이가 A4 1장을 넘지 않도록 조절""",
        max_tokens=2500,
    )
    # HTML 추출
    import re as _re
    m = _re.search(r'<!DOCTYPE.*?</html>|<html.*?</html>', resp, _re.DOTALL | _re.IGNORECASE)
    return m.group() if m else resp


def worksheet_to_pdf(html: str) -> bytes | None:
    """HTML → PDF 바이트 (weasyprint)"""
    try:
        from weasyprint import HTML as WH
        return WH(string=html).write_pdf()
    except Exception as e:
        return None


# ── 학생 활동지 생성 ──────────────────────────────────────────────
def gen_worksheet(grade, theme, book, book_info, questions: dict = None,
                  book_ctx_extra: str = "") -> str:
    """AI가 A4 HTML 활동지 생성"""
    pub    = book_info.get("publisher","") if book_info else ""
    author = book_info.get("author","")    if book_info else ""
    book_str = f"{author} · {pub}" if pub else author

    # 질문 카드에서 샘플 추출
    q_lines = ""
    if questions:
        all_qs = []
        for sec in ["before","during","after"]:
            for q in questions.get(sec,[]):
                all_qs.append(q.get("text","") if isinstance(q,dict) else str(q))
        q_lines = "\n".join(f"- {q}" for q in all_qs[:6])

    return chat(
        "초등 그림책 수업 활동지 디자이너입니다. 완전한 HTML만 출력하고 다른 텍스트는 절대 쓰지 마세요.",
        f"""그림책: 「{book}」 ({book_str}) / 학년: {grade} / 주제: {theme}

참고 질문:
{q_lines}

아래 조건을 모두 지켜 학생 활동지 HTML을 만드세요.

[디자인]
- Google Fonts: Nanum Gothic + Gaegu (CDN)
- A4 1장 (210×297mm), @media print 포함
- -webkit-print-color-adjust: exact
- 파스텔 배경 (#FFF9F0 또는 유사), 컬러풀한 섹션 헤더
- 이모지 적극 사용, 점선 답변칸

[구성 — 정확히 이 순서로]
1. 상단 헤더: 그림책 제목(Gaegu 큰 글씨) + 이름·날짜란
2. 읽기 전 📖: 예측/상상 질문 1개 + 답변칸(점선 3줄)
3. 읽으면서 🔍: 핵심 질문 2개 + 각 답변칸(점선 2줄)
4. 읽은 후 💬: 추론/삶연결 질문 2개 + 각 답변칸(점선 2줄)
5. 나만의 표현 🎨: 빈 박스 (그림/글 자유 표현) — 전체 높이의 약 20%
6. 하단: "오늘의 한마디" 한 줄 답변칸

모든 질문은 참고 질문에서 가져오되 학년 수준에 맞게 다듬으세요.
반드시 <!DOCTYPE html>부터 </html>까지 완전한 HTML로만 답하세요.""",
        max_tokens=2500,
    )


def worksheet_to_pdf(html: str) -> bytes | None:
    """HTML → PDF bytes (weasyprint)"""
    try:
        from weasyprint import HTML as WH
        return WH(string=html, base_url=None).write_pdf()
    except Exception:
        return None

# ── PDF 그림책 내용 추출 ──────────────────────────────────────────
def extract_pdf_text(uploaded_file) -> str:
    """업로드된 PDF에서 텍스트 추출 (PyMuPDF 우선, 없으면 pdfplumber)"""
    try:
        import fitz  # PyMuPDF
        data = uploaded_file.read()
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text())
        text = "\n".join(pages).strip()
        if len(text) > 100:
            return text[:4000]  # 토큰 절약
        # 텍스트가 너무 적으면 스캔본 → Vision API로 OCR
        return _ocr_pdf_with_vision(data, doc)
    except ImportError:
        pass
    try:
        import pdfplumber
        from io import BytesIO
        uploaded_file.seek(0)
        with pdfplumber.open(BytesIO(uploaded_file.read())) as pdf:
            text = "\n".join(p.extract_text() or "" for p in pdf.pages[:10])
        return text.strip()[:4000] if text.strip() else ""
    except Exception as e:
        return ""

def _ocr_pdf_with_vision(pdf_bytes: bytes, doc) -> str:
    """스캔 PDF: 첫 3페이지를 이미지로 변환 후 Vision API OCR"""
    import base64
    client = get_client()
    results = []
    for i, page in enumerate(doc):
        if i >= 3:
            break
        pix = page.get_pixmap(dpi=150)
        img_b64 = base64.b64encode(pix.tobytes("png")).decode()
        try:
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image_url",
                         "image_url": {"url": f"data:image/png;base64,{img_b64}"}},
                        {"type": "text",
                         "text": "이 그림책 페이지의 텍스트를 그대로 읽어주세요. 글이 없으면 그림 내용을 간략히 묘사해 주세요."}
                    ]
                }],
                max_tokens=400,
            )
            results.append(resp.choices[0].message.content or "")
        except Exception:
            pass
    return "\n".join(results)[:4000]

def summarize_book_content(raw_text: str, book_title: str) -> str:
    """추출된 텍스트를 수업용 줄거리·특성으로 요약"""
    return chat(
        "초등 그림책 전문가입니다. 한국어로 답합니다.",
        f"""아래는 그림책 「{book_title}」의 내용입니다.
다음 형식으로 요약해 주세요.

[줄거리] 3~4문장
[주요 등장인물] 간략히
[핵심 메시지] 1~2문장
[수업 활용 포인트] 초등 수업에서 어떻게 활용할 수 있는지 2~3가지

---
{raw_text[:3000]}""",
        max_tokens=600,
    )

# ── 웹 검색으로 그림책 정보 조회 ─────────────────────────────────
def search_book_online(book_title: str) -> str:
    """OpenAI로 그림책 정보 검색 (웹 검색 대체 — 모델 지식 활용)"""
    return chat(
        "초등 그림책 전문가입니다. 한국어로 답합니다.",
        f"""그림책 「{book_title}」에 대해 알려주세요.
모르는 책이라면 솔직히 말해주세요.

[작가/출판사]
[줄거리] 3~4문장
[주요 등장인물]
[핵심 메시지] 1~2문장
[초기 문해력 요소] 음운인식/어휘/이야기이해/추론/배경지식/감정이해 중 해당하는 것
[수업 활용 포인트] 2~3가지""",
        max_tokens=600,
    )

# ── DOCX ─────────────────────────────────────────────────────────
def make_docx(sections: dict, title: str) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for sec_title, content in sections.items():
        if content:
            doc.add_heading(sec_title, level=2)
            for line in content.splitlines():
                s = line.strip()
                if not s: doc.add_paragraph("")
                elif s.startswith("## "): doc.add_heading(s[3:], level=3)
                elif s.startswith(("- ","• ")):
                    p = doc.add_paragraph(style="List Bullet"); p.add_run(s[2:])
                else: doc.add_paragraph(s)
    buf = BytesIO(); doc.save(buf); return buf.getvalue()

# ── PPTX ─────────────────────────────────────────────────────────
def _parse_evals(raw: str) -> list:
    """eval_parent JSON에서 criteria 추출"""
    if not raw: return []
    import re
    try:
        m = re.search(r'\{.*\}', raw, re.DOTALL)
        data = json.loads(m.group()) if m else {}
        return data.get("criteria", [])
    except Exception:
        return []

def make_pptx(grade, theme, book, lesson_time, questions, activities_text, eval_raw: str = "") -> bytes | None:
    # 활동 파싱 (JSON 구조)
    acts = []
    if activities_text:
        import re
        try:
            m = re.search(r'\[.*\]', activities_text, re.DOTALL)
            act_list = json.loads(m.group()) if m else []
            icons = ["🌱","📖","✍️"]
            for i, a in enumerate(act_list):
                acts.append({
                    "icon": icons[i % 3],
                    "title": a.get("title",""),
                    "desc": " ".join(a.get("process",[][:1])) or a.get("goal","")
                })
        except Exception:
            pass

    # 수업 목표 (간단 생성)
    obj_resp = chat(
        "한국어로 간결하게 답변하세요.",
        f"그림책 '{book}'로 {grade} {theme} 수업 목표 3개를 각각 한 문장으로 JSON 배열로만 답하세요. [\"목표1\",\"목표2\",\"목표3\"]",
        max_tokens=200,
    )
    try:
        import re
        m = re.search(r'\[.*?\]', obj_resp, re.DOTALL)
        objectives = json.loads(m.group()) if m else ["목표를 설정합니다.","질문을 만들어 봅니다.","생각을 표현합니다."]
    except Exception:
        objectives = ["목표를 설정합니다.","질문을 만들어 봅니다.","생각을 표현합니다."]

    data = {
        "title": f"「{book}」 질문수업",
        "subtitle": "AI 그림책 질문수업 설계안",
        "grade": grade,
        "theme": theme,
        "book": book,
        "lesson_time": lesson_time,
        "objectives": objectives,
        "questions": questions or {},
        "activities": acts or [
            {"icon":"🌱","title":"배경지식 활성화","desc":"그림책 표지 탐색, 경험 나누기"},
            {"icon":"📖","title":"대화형 읽기","desc":"PEER 절차로 그림책 읽기, 질문-응답"},
            {"icon":"✍️","title":"표현 활동","desc":"느낀 점 쓰기, 그림으로 표현하기"},
        ],
        "evaluations": _parse_evals(eval_raw),
    }

    json_str = json.dumps(data, ensure_ascii=False)
    with tempfile.NamedTemporaryFile(suffix=".pptx", delete=False) as f:
        out_path = f.name

    try:
        result = subprocess.run(
            ["node", PPTX_SCRIPT, json_str, out_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            st.error(f"PPT 생성 오류: {result.stderr[:200]}")
            return None
        with open(out_path, "rb") as f:
            return f.read()
    except Exception as e:
        st.error(f"PPT 생성 실패: {e}")
        return None
    finally:
        try: os.unlink(out_path)
        except: pass

# ═══════════════════════════════════════════════════════════════════
# ── 결과 섹션 렌더러 (expander 완전 대체) ───────────────────────
def result_section(label: str, content_fn, *args, **kwargs):
    """헤더 + 내용을 하나의 styled div로 렌더링"""
    st.markdown(
        f'<div class="result-section">',
        unsafe_allow_html=True
    )
    st.markdown(
        f'<div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;' +
        f'color:#111827;">{label}</div>',
        unsafe_allow_html=True
    )
    content_fn(*args, **kwargs)
    st.markdown('</div>', unsafe_allow_html=True)

# ── 지도안 테이블 렌더러 ─────────────────────────────────────────
def render_lessonplan(raw: str):
    """gen_lessonplan JSON → 상세 테이블 UI (질문 번호, 예상 답변 포함)"""
    import re as _re
    try:
        m = _re.search(r'\[.*\]', raw, _re.DOTALL)
        stages = json.loads(m.group()) if m else []
    except Exception:
        st.markdown(raw)
        return

    stage_cfg = {
        "도입": ("#3B82F6", "#EFF6FF"),
        "전개": ("#8B5CF6", "#F5F3FF"),
        "정리": ("#10B981", "#F0FDF4"),
    }

    def fmt_teacher(row: dict, q_color: str) -> str:
        """질문 번호 + 줄바꿈 + 교사 활동 텍스트"""
        qs = row.get("questions", [])
        activity = row.get("activity", "")
        teacher  = row.get("teacher", "")
        parts = []
        if activity:
            parts.append(f'<div style="font-weight:700;font-size:.8rem;color:{q_color};margin-bottom:4px;">▶ {activity}</div>')
        if qs:
            q_lines = "".join(
                f'<div style="display:flex;gap:5px;margin-bottom:4px;">' +
                f'<span style="min-width:18px;font-weight:700;color:{q_color};font-size:.78rem;">{i+1}.</span>' +
                f'<span style="font-size:.82rem;color:#1D4ED8;line-height:1.5;">{q}</span></div>'
                for i, q in enumerate(qs)
            )
            parts.append(f'<div style="background:#EFF6FF;border-radius:6px;padding:5px 7px;margin-bottom:5px;">{q_lines}</div>')
        if teacher:
            parts.append(f'<div style="font-size:.82rem;color:#374151;line-height:1.5;">{teacher}</div>')
        return "".join(parts)

    def fmt_student(row: dict) -> str:
        """학생 활동 + 예상 답변을 Q번호별로 구분"""
        student = row.get("student", "")
        if not student:
            return ""
                # "A1:" 또는 "A1 —" 패턴으로 분리 (학생 예상 답변)
        segments = _re.split(r'A(\d+)[\s]*[—\-:]', student)
        if len(segments) <= 1:
            return f'<div style="font-size:.82rem;color:#166534;line-height:1.6;">{student}</div>'
        result = f'<div style="font-size:.82rem;color:#166534;line-height:1.5;">{segments[0]}</div>' if segments[0].strip() else ""
        for j in range(1, len(segments), 2):
            anum = segments[j]
            content = segments[j+1].strip() if j+1 < len(segments) else ""
            result += (
                f'<div style="display:flex;gap:5px;margin:3px 0;padding:4px 7px;' +
                f'background:#F0FDF4;border-radius:5px;">' +
                f'<span style="min-width:22px;font-weight:700;color:#059669;font-size:.75rem;">A{anum}</span>' +
                f'<span style="font-size:.81rem;color:#166534;line-height:1.5;">{content}</span></div>'
            )
        return result

    header = '''<table class="lp-table">
      <thead><tr>
        <th style="width:6%;">단계</th>
        <th style="width:6%;">시간</th>
        <th style="width:35%;">🧑‍🏫 교사 발문 · 활동</th>
        <th style="width:35%;">🙋 학생 활동 · 예상 반응</th>
        <th style="width:18%;">💡 유의점</th>
      </tr></thead><tbody>'''
    rows_html = ""

    for stage in stages:
        name  = stage.get("stage", "")
        dur   = stage.get("duration", "")
        rows  = stage.get("rows", [])
        color, bg = stage_cfg.get(name, ("#6B7280", "#F9FAFB"))
        rcount = len(rows)

        for i, row in enumerate(rows):
            teacher_html = fmt_teacher(row, color)
            student_html = fmt_student(row)
            note = row.get("note", "")
            note_html = f'<div class="lp-note" style="font-size:.78rem;line-height:1.5;">{note}</div>' if note else ""

            if i == 0:
                rows_html += f"""<tr>
  <td rowspan="{rcount}" style="text-align:center;background:{bg};border-color:#E5E7EB;">
    <span class="lp-stage" style="background:{color};color:white;">{name}</span>
  </td>
  <td rowspan="{rcount}" class="lp-dur">{dur}</td>
  <td class="lp-teacher" style="padding:10px 12px;">{teacher_html}</td>
  <td class="lp-student" style="padding:10px 12px;">{student_html}</td>
  <td style="padding:10px 12px;">{note_html}</td>
</tr>"""
            else:
                rows_html += f"""<tr>
  <td class="lp-teacher" style="padding:10px 12px;">{teacher_html}</td>
  <td class="lp-student" style="padding:10px 12px;">{student_html}</td>
  <td style="padding:10px 12px;">{note_html}</td>
</tr>"""

    footer = "</tbody></table>"
    st.markdown(header + rows_html + footer, unsafe_allow_html=True)


# ── 활동 카드 렌더러 ─────────────────────────────────────────────
def render_activities(raw: str):
    """gen_activities JSON → 카드 UI"""
    import re as _re
    try:
        m = _re.search(r'\[.*\]', raw, _re.DOTALL)
        acts = json.loads(m.group()) if m else []
    except Exception:
        st.markdown(raw)
        return
    step_colors = {"도입": "#3B82F6", "중심": "#8B5CF6", "정리": "#10B981"}
    for act in acts:
        color = step_colors.get(act.get("step",""), "#1A1A2E")
        process_html = "".join(
            f'<div style="display:flex;gap:6px;margin-bottom:3px;">'
            f'<span style="min-width:18px;font-weight:700;color:{color};">{i+1}.</span>'
            f'<span style="font-size:.82rem;color:#374151;line-height:1.5;">{p}</span></div>'
            for i, p in enumerate(act.get("process", []))
        )
        st.markdown(f"""
<div class="act-card">
  <div class="act-header">
    <span class="act-step-badge" style="background:{color};">{act.get('step','')}</span>
    <span class="act-title">{act.get('title','')}</span>
    <span class="act-duration">⏱ {act.get('duration','')}</span>
  </div>
  <div class="act-row">
    <span class="act-label">🎯 목표</span>
    <span class="act-val">{act.get('goal','')}</span>
  </div>
  <div class="act-row">
    <span class="act-label">📦 준비물</span>
    <span class="act-val">{act.get('materials','')}</span>
  </div>
  <div class="act-row">
    <span class="act-label">📋 진행</span>
    <div>{process_html}</div>
  </div>
  <div class="act-question">💬 "{act.get('question','')}"</div>
</div>""", unsafe_allow_html=True)


# ── 평가 카드 렌더러 ──────────────────────────────────────────────
def render_eval_parent(raw: str):
    """gen_eval_parent JSON → 카드 UI"""
    import re as _re
    try:
        m = _re.search(r'\{.*\}', raw, _re.DOTALL)
        data = json.loads(m.group()) if m else {}
    except Exception:
        st.markdown(raw)
        return

    criteria = data.get("criteria", [])
    self_eval = data.get("self_eval", [])
    parent_letter = data.get("parent_letter", "")

    if criteria:
        st.markdown("#### 📊 관찰 평가 기준")
        rows = "".join(
            f'''<tr>
              <td class="td-area">{c.get("area","")}</td>
              <td style="color:#166534;">{c.get("good","")}</td>
              <td style="color:#92400E;">{c.get("ok","")}</td>
              <td style="color:#1D4ED8;">{c.get("needs","")}</td>
            </tr>'''
            for c in criteria
        )
        st.markdown(f"""
<table class="eval-table">
  <thead><tr>
    <th class="th-area">평가 영역</th>
    <th class="th-good">✅ 잘함</th>
    <th class="th-ok">🟡 보통</th>
    <th class="th-need">🔵 노력 필요</th>
  </tr></thead>
  <tbody>{rows}</tbody>
</table>""", unsafe_allow_html=True)

    if self_eval:
        st.markdown("#### ✏️ 학생 자기평가")
        for i, q in enumerate(self_eval):
            st.markdown(
                f'<div style="background:#F9FAFB;border:1px solid #E5E7EB;border-radius:8px;'
                f'padding:.55rem .9rem;margin-bottom:6px;font-size:.85rem;color:#374151;"'
                f'><span style="font-weight:700;color:#6B7280;margin-right:8px;">{i+1}</span>{q}</div>',
                unsafe_allow_html=True
            )

    if parent_letter:
        st.markdown("#### 👨‍👩‍👧 학부모 안내문")
        st.markdown(
            f'<div style="background:#FFF9F0;border:1px solid #FDE68A;border-radius:10px;'
            f'padding:1rem 1.1rem;font-size:.85rem;color:#374151;line-height:1.8;white-space:pre-wrap;"'
            f'>{parent_letter}</div>',
            unsafe_allow_html=True
        )


# 질문 카드 렌더러
# ═══════════════════════════════════════════════════════════════════
def render_question_cards(questions: dict):
    sections = [
        ("before", "🌱 읽기 전"),
        ("during", "🔍 읽는 중"),
        ("after",  "💬 읽은 후"),
    ]
    for key, label in sections:
        qs = questions.get(key, [])
        if not qs: continue
        st.markdown(f"**{label}**")
        cols = st.columns(min(len(qs), 3))
        for i, q in enumerate(qs):
            qtype = q.get("type", "")
            qtext = q.get("text", q) if isinstance(q, dict) else q
            with cols[i % len(cols)]:
                st.markdown(
                    f'<div class="q-card">'
                    f'<span class="qt {qtype}">{qtype}</span>'
                    f'<div class="qtext">{qtext}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )
        st.markdown("<br>", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════
def main():
    st.markdown(CSS, unsafe_allow_html=True)

    # ── 헤더 ──
    st.markdown("""
    <div class="top-bar">
      <span class="top-bar-icon">📚</span>
      <span class="top-bar-title">AI 그림책 질문수업 설계기</span>
      <span class="top-bar-sub">초기 문해력 수업 코파일럿</span>
    </div>
    """, unsafe_allow_html=True)

    # ── STEP 1: 수업 조건 ──────────────────────────────────────────
    st.markdown('<div class="section-title st-blue">수업 조건</div>', unsafe_allow_html=True)
    # 라벨을 selectbox 위에 표시
    lc1, lc2, lc3 = st.columns(3)
    with lc1: st.caption("📌 학년")
    with lc2: st.caption("🎯 주제")
    with lc3: st.caption("⏰ 수업 시간")
    c1, c2, c3 = st.columns(3)
    with c1: grade = st.selectbox("학년", ["초등 1학년","초등 2학년","초등 3학년",
                                            "초등 4학년","초등 5학년","초등 6학년"],
                                   label_visibility="collapsed")
    with c2: theme = st.selectbox("주제", db_all_themes(), label_visibility="collapsed")
    with c3: lesson_time = st.selectbox("시간", ["40분","80분","120분","3차시","5차시"],
                                         label_visibility="collapsed")

    st.caption("📝 학생 특성 (선택)")
    student_context = st.text_area("학생 특성 (선택)",
        placeholder="예: 1학년 입학 초기, 친구 관계가 서툰 편, 글쓰기 부담이 큼 등",
        height=64, label_visibility="collapsed")

    # ── STEP 2: 그림책 선택 ────────────────────────────────────────
    st.markdown('<div class="section-title st-purple">그림책 선택</div>', unsafe_allow_html=True)

    book_tab1, book_tab2, book_tab3 = st.tabs(["🤖 AI 추천", "📚 DB에서 찾기", "✏️ 직접 입력"])

    # 탭별로 독립적인 book/book_info 관리
    # 각 탭은 자신이 활성일 때만 book/book_info를 session_state에 저장
    book = ""
    book_info = None
    custom_summary = ""
    input_method = st.session_state.get("input_method", "📝 제목 입력")

    # ─ AI 추천 탭 ─
    with book_tab1:
        situation_input = st.text_area(
            "우리 반 상황 입력",
            placeholder='예: "우리 반은 친구 관계 갈등이 많아요"\n"자존감이 낮아 자기 표현을 못하는 아이들이 있어요"',
            height=80, key="ai_situation", label_visibility="collapsed"
        )
        st.caption("📌 우리 반 상황을 자유롭게 입력하면 AI가 맞는 그림책을 추천해 드려요")

        if st.button("🤖 그림책 추천받기", key="btn_ai_rec"):
            if not situation_input.strip():
                st.warning("상황을 입력해 주세요.")
            else:
                with st.spinner("AI가 그림책을 고르는 중..."):
                    recs = ai_recommend_books(situation_input)
                st.session_state["ai_recs"] = recs

        if "ai_recs" in st.session_state and st.session_state["ai_recs"]:
            st.markdown('<div style="font-size:.8rem;font-weight:600;color:#6B7280;margin-bottom:.5rem;">추천 그림책 — 버튼을 누르면 바로 적용됩니다</div>', unsafe_allow_html=True)
            recs = st.session_state["ai_recs"]
            cols = st.columns(len(recs))
            for i, (r, col) in enumerate(zip(recs, cols)):
                is_selected = st.session_state.get("active_book") == r["title"] and st.session_state.get("active_tab") == "ai"
                label = f"{'✅ ' if is_selected else ''}{i+1}. {r['title']}"
                with col:
                    if st.button(label, key=f"btn_ai_book_{i}", use_container_width=True):
                        st.session_state["active_tab"] = "ai"
                        st.session_state["active_book"] = r["title"]
                        st.session_state["active_book_info"] = db_get_by_title(r["title"])
                        st.session_state.pop("active_custom_summary", None)
                        st.rerun()
            # 선택된 책의 추천 이유 표시
            active = st.session_state.get("active_book")
            for r in recs:
                if r["title"] == active and r.get("reason"):
                    st.caption(f"💡 {r['reason']}")
                    break

    # ─ DB 탭 ─
    with book_tab2:
        show_db = st.toggle("🔍 전체 DB 탐색 펼치기", value=False)
        if show_db:
            fc1, fc2 = st.columns(2)
            with fc1: ft = st.selectbox("주제", ["전체"] + db_all_themes(), key="dbt")
            with fc2: fg = st.selectbox("학년", ["전체","초등 1학년","초등 2학년",
                                                   "초등 3학년","초등 4학년","초등 5학년","초등 6학년"], key="dbg")
            filtered = db_search("" if ft=="전체" else ft, "" if fg=="전체" else fg)
            st.caption(f"검색 결과 {len(filtered)}권")
            cards = ""
            for b in filtered:
                tags = "".join(f'<span class="tag">{t}</span>' for t in b["theme"][:2])
                pub_b = b.get("publisher","")
                pub_str_b = f' · {pub_b}' if pub_b else ''
                cards += (f'<div class="book-card" style="margin-bottom:6px;">'
                          f'<span class="bc-icon">📕</span>'
                          f'<div><div class="bc-title">{b["title"]}</div>'
                          f'<div class="bc-meta">{b["author"]}{pub_str_b}</div>'
                          f'<div class="bc-tags">{tags}</div></div></div>')
            st.markdown(cards, unsafe_allow_html=True)

        rec = db_search(theme=theme, grade=grade)
        if rec:
            sel = st.selectbox(f"추천 ({len(rec)}권 — {grade} × {theme})",
                               [b["title"] for b in rec])
            # DB 탭이 포커스될 때만 active_tab을 "db"로 설정
            if st.button("✅ 이 책으로 선택", key="btn_db_select", use_container_width=False):
                st.session_state["active_tab"] = "db"
                st.session_state["active_book"] = sel
                st.session_state["active_book_info"] = db_get_by_title(sel)
                st.session_state.pop("active_custom_summary", None)
                st.rerun()
            # 현재 active_tab이 db인 경우에만 book 설정
            if st.session_state.get("active_tab") == "db":
                saved = st.session_state.get("active_book", "")
                if saved in [b["title"] for b in rec]:
                    book = saved
                    book_info = st.session_state.get("active_book_info")
        else:
            st.info("조건에 맞는 책이 없습니다.")

    # ─ 직접 입력 탭 ─
    with book_tab3:
        # 세 가지 방식 선택
        input_method = st.radio(
            "입력 방식",
            ["📝 제목 입력", "📄 PDF 업로드", "🌐 웹 검색"],
            horizontal=True,
            key="input_method",
            label_visibility="collapsed",
        )
        st.caption("📝 제목 직접 입력  　📄 그림책 PDF 업로드  　🌐 책 정보 웹 검색")

        custom_summary = ""  # PDF/웹 검색으로 얻은 추가 정보

        # ── 📝 제목 입력 ──
        if input_method == "📝 제목 입력":
            custom = st.text_input("그림책 제목", placeholder="예: 알사탕", key="custom_book_title")
            if custom:
                _bi = db_get_by_title(custom)
                if _bi:
                    st.success("✅ DB에 있는 책입니다!")
                else:
                    st.info("ℹ️ DB에 없는 책 — AI 일반 지식으로 진행합니다.")
                if st.button("✅ 이 책으로 선택", key="btn_title_select"):
                    st.session_state["active_tab"] = "title"
                    st.session_state["active_book"] = custom
                    st.session_state["active_book_info"] = _bi
                    st.session_state.pop("active_custom_summary", None)
                    st.rerun()

        # ── 📄 PDF 업로드 ──
        elif input_method == "📄 PDF 업로드":
            st.caption("그림책을 스캔한 PDF 또는 텍스트 PDF를 올려주세요.")
            pdf_file = st.file_uploader(
                "PDF 업로드", type=["pdf"], key="pdf_upload",
                label_visibility="collapsed"
            )
            pdf_title = st.text_input("그림책 제목 (필수)", placeholder="예: 우리 선생님이 최고야", key="pdf_title")

            if pdf_file and pdf_title:
                if st.button("📖 PDF 내용 분석하기", key="btn_pdf"):
                    with st.spinner("PDF를 읽는 중... (스캔본은 조금 더 걸릴 수 있어요)"):
                        raw = extract_pdf_text(pdf_file)
                    if raw:
                        with st.spinner("그림책 내용을 수업용으로 요약 중..."):
                            summary = summarize_book_content(raw, pdf_title)
                        st.session_state["custom_summary"] = summary
                        st.session_state["custom_title"] = pdf_title
                        st.success("✅ 분석 완료!")
                    else:
                        st.error("텍스트를 추출할 수 없었습니다. 파일을 확인해 주세요.")

                if "custom_summary" in st.session_state and st.session_state.get("custom_title") == pdf_title:
                    st.session_state["active_tab"] = "pdf"
                    st.session_state["active_book"] = pdf_title
                    st.session_state["active_book_info"] = None
                    book = pdf_title
                    custom_summary = st.session_state["custom_summary"]
                    st.session_state["active_custom_summary"] = custom_summary
                    st.markdown('<div class="result-section"><div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;color:#111827;">📋  분석된 내용 확인</div>', unsafe_allow_html=True)
                    st.markdown(custom_summary)
                    st.markdown('</div>', unsafe_allow_html=True)

        # ── 🌐 웹 검색 ──
        elif input_method == "🌐 웹 검색":
            st.caption("책 제목을 입력하면 AI가 책 정보를 검색해 드려요.")
            search_title = st.text_input(
                "그림책 제목 검색", placeholder="예: 100만 번 산 고양이",
                key="web_search_title"
            )
            if search_title:
                if st.button("🔍 책 정보 검색", key="btn_web_search"):
                    with st.spinner(f"「{search_title}」 정보를 찾는 중..."):
                        info = search_book_online(search_title)
                    st.session_state["web_search_result"] = info
                    st.session_state["web_search_title"] = search_title

                if (st.session_state.get("web_search_title") == search_title
                        and "web_search_result" in st.session_state):
                    st.session_state["active_tab"] = "web"
                    st.session_state["active_book"] = search_title
                    st.session_state["active_book_info"] = None
                    book = search_title
                    custom_summary = st.session_state["web_search_result"]
                    st.session_state["active_custom_summary"] = custom_summary
                    st.markdown('<div class="result-section"><div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;color:#111827;">📋  검색된 책 정보</div>', unsafe_allow_html=True)
                    st.markdown(custom_summary)
                    st.markdown('</div>', unsafe_allow_html=True)
                    st.success("✅ 이 정보로 수업안을 만듭니다.")

    # ── 최종 book/book_info 결정: active_tab 기준 ──────────────────
    # 탭과 무관하게 session_state의 active_tab이 기준
    active_tab = st.session_state.get("active_tab", "")

    if active_tab in ("ai", "db", "title"):
        book = st.session_state.get("active_book", "")
        book_info = st.session_state.get("active_book_info")
        custom_summary = ""
    elif active_tab in ("pdf", "web"):
        book = st.session_state.get("active_book", "")
        book_info = None
        custom_summary = st.session_state.get("active_custom_summary", "")
    # active_tab이 없으면 book=""로 유지 (초기 상태)

    # 선택된 책 표시
    if book:
        if book_info:
            tags_html = "".join(f'<span class="tag">{e}</span>' for e in book_info["literacy_elements"])
            pub = book_info.get("publisher", "")
            pub_str = f' · {pub}' if pub else ''
            st.markdown(
                f'<div class="book-card">'
                f'<span class="bc-icon">📕</span>'
                f'<div><div class="bc-title">{book_info["title"]}</div>'
                f'<div class="bc-meta">{book_info["author"]}{pub_str}</div>'
                f'<div class="bc-meta" style="color:#9CA3AF;font-size:.72rem;">{book_info["summary"]}</div>'
                f'<div class="bc-tags">{tags_html}</div></div></div>',
                unsafe_allow_html=True)
        elif st.session_state.get("active_custom_summary"):
            st.markdown(
                f'<div class="book-card"><span class="bc-icon">📕</span>'
                f'<div><div class="bc-title">{book}</div>'
                f'<div class="bc-meta" style="color:#5D8A65;">'
                f'✅ 내용 분석 완료 — 이 정보로 수업안을 생성합니다</div></div></div>',
                unsafe_allow_html=True)
        else:
            st.info(f"📖 선택된 책: **{book}**")


    # ── STEP 3: 수업안 생성 ─────────────────────────────────────────
    st.markdown('<div class="section-title st-green">수업안 생성</div>', unsafe_allow_html=True)

    if not book:
        st.info("책을 먼저 선택해 주세요.")
    else:
        _ctx_extra = st.session_state.get("active_custom_summary", "")

        # ── 완료 현황 배지 ──────────────────────────────────────────
        steps = [
            ("questions",   "❓ 질문"),
            ("activities",  "🎨 활동"),
            ("lessonplan",  "🗒️ 지도안"),
            ("eval_parent", "⭐ 평가"),
            ("worksheet",   "📋 활동지"),
        ]
        done_count = sum(1 for k, _ in steps if k in st.session_state)
        status_parts = []
        for k, label in steps:
            if k in st.session_state:
                status_parts.append(f'<span style="background:#F0FDF4;color:#166534;border:1px solid #BBF7D0;border-radius:4px;padding:2px 8px;font-size:.75rem;font-weight:600;">{label} ✓</span>')
            else:
                status_parts.append(f'<span style="background:#F9FAFB;color:#9CA3AF;border:1px solid #E5E7EB;border-radius:4px;padding:2px 8px;font-size:.75rem;">{label}</span>')
        st.markdown(
            '<div style="display:flex;gap:6px;flex-wrap:wrap;margin-bottom:.8rem;">' +
            " ".join(status_parts) + "</div>",
            unsafe_allow_html=True
        )

        # ── 메인: 전체 생성 버튼 ───────────────────────────────────
        if st.button("✨ 수업안 전체 생성 (질문 + 활동 + 지도안 + 평가)", type="primary", use_container_width=True, key="btn_all"):
            prog = st.progress(0, text="수업안을 만드는 중...")
            with st.spinner("❓ 질문 생성 중..."):
                qs = gen_questions(grade, theme, book, book_info, _ctx_extra)
                st.session_state["questions"] = qs
            prog.progress(25, text="🎨 활동 생성 중...")
            acts = gen_activities(grade, theme, book, lesson_time, student_context, _ctx_extra)
            st.session_state["activities"] = acts
            prog.progress(50, text="🗒️ 지도안 생성 중...")
            lp = gen_lessonplan(grade, theme, book, lesson_time, student_context, _ctx_extra,
                                       questions=st.session_state.get('questions', {}))
            st.session_state["lessonplan"] = lp
            prog.progress(75, text="⭐ 평가·안내문 생성 중...")
            ev = gen_eval_parent(grade, theme, book, _ctx_extra)
            st.session_state["eval_parent"] = ev
            prog.progress(100, text="완료!")
            st.rerun()

        # ── 보조: 개별 재생성 ──────────────────────────────────────
        if st.toggle("🔄 개별 항목 재생성", key="toggle_regen"):
            sc1, sc2, sc3, sc4 = st.columns(4)
            with sc1:
                if st.button("❓ 질문 재생성", use_container_width=True, key="btn_q"):
                    with st.spinner(""):
                        st.session_state["questions"] = gen_questions(grade, theme, book, book_info, _ctx_extra)
            with sc2:
                if st.button("🎨 활동 재생성", use_container_width=True, key="btn_a"):
                    with st.spinner(""):
                        st.session_state["activities"] = gen_activities(grade, theme, book, lesson_time, student_context, _ctx_extra)
            with sc3:
                if st.button("🗒️ 지도안 재생성", use_container_width=True, key="btn_l"):
                    with st.spinner(""):
                        st.session_state["lessonplan"] = gen_lessonplan(grade, theme, book, lesson_time, student_context, _ctx_extra,
                                                                                  questions=st.session_state.get('questions', {}))
            with sc4:
                if st.button("⭐ 평가 재생성", use_container_width=True, key="btn_e"):
                    with st.spinner(""):
                        st.session_state["eval_parent"] = gen_eval_parent(grade, theme, book, _ctx_extra)

        # ── 결과 ──────────────────────────────────────────────────
        has_any = any(k in st.session_state for k in
                      ["questions","activities","lessonplan","eval_parent","worksheet"])
        if has_any:
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.markdown('<div class="section-title">생성 결과</div>', unsafe_allow_html=True)

            if "questions" in st.session_state and st.session_state["questions"]:
                st.markdown('<div class="result-section"><div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;color:#111827;">❓  질문 카드</div>', unsafe_allow_html=True)
                render_question_cards(st.session_state["questions"])
                st.markdown('</div>', unsafe_allow_html=True)

            if "activities" in st.session_state:
                st.markdown('<div class="result-section"><div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;color:#111827;">🎨  활동 생성</div>', unsafe_allow_html=True)
                render_activities(st.session_state["activities"])
                st.markdown('</div>', unsafe_allow_html=True)

            if "lessonplan" in st.session_state:
                st.markdown('<div class="result-section"><div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;color:#111827;">🗒️  지도안</div>', unsafe_allow_html=True)
                render_lessonplan(st.session_state["lessonplan"])
                st.markdown('</div>', unsafe_allow_html=True)

            if "eval_parent" in st.session_state:
                st.markdown('<div class="result-section"><div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;color:#111827;">⭐  평가 & 학부모 안내문</div>', unsafe_allow_html=True)
                render_eval_parent(st.session_state["eval_parent"])
                st.markdown('</div>', unsafe_allow_html=True)

            # ── 학생 활동지 ─────────────────────────────────────────
            st.markdown('<div class="result-section"><div style="padding:.65rem 1rem .5rem;font-size:.9rem;font-weight:600;color:#111827;">📋  학생 활동지</div>', unsafe_allow_html=True)
            if "worksheet" in st.session_state:
                # HTML 미리보기
                ws_html = st.session_state["worksheet"]
                st.components.v1.html(ws_html, height=900, scrolling=True)
                st.markdown("<br>", unsafe_allow_html=True)
                # PDF 다운로드
                ws_pdf = worksheet_to_pdf(ws_html)
                if ws_pdf:
                    st.download_button(
                        "⬇️ 활동지 PDF 다운로드",
                        data=ws_pdf,
                        file_name=f"{book}_{grade}_활동지.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                    )
                else:
                    st.download_button(
                        "⬇️ 활동지 HTML 다운로드",
                        data=ws_html.encode("utf-8"),
                        file_name=f"{book}_{grade}_활동지.html",
                        mime="text/html",
                        use_container_width=True,
                    )
            else:
                has_q = "questions" in st.session_state
                if st.button(
                    "📋 활동지 생성하기" if has_q else "📋 활동지 생성 (질문 먼저 생성해 주세요)",
                    use_container_width=True,
                    key="btn_ws",
                    disabled=not has_q,
                ):
                    with st.spinner("✏️ 활동지를 디자인하는 중..."):
                        ws = gen_worksheet(
                            grade, theme, book, book_info,
                            questions=st.session_state.get("questions", {}),
                            book_ctx_extra=_ctx_extra,
                        )
                    st.session_state["worksheet"] = ws
                    st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

            # ── 다운로드 & PPT ───────────────────────────────────
            st.markdown("<br>", unsafe_allow_html=True)
            dl1, dl2, dl3 = st.columns([2, 2, 3])

            sections_for_docx = {
                "질문 생성": str(st.session_state.get("questions","")),
                "활동": st.session_state.get("activities",""),
                "지도안": st.session_state.get("lessonplan",""),
                "평가·학부모 안내문": st.session_state.get("eval_parent",""),
            }
            docx_title = f"{book}_{theme}_질문수업설계안"

            with dl1:
                st.download_button(
                    "📄 Word",
                    data=make_docx(sections_for_docx, docx_title),
                    file_name=f"{docx_title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with dl2:
                combined_md = "\n\n".join(
                    f"## {k}\n{v}" for k, v in sections_for_docx.items() if v
                )
                st.download_button(
                    "📝 Markdown",
                    data=combined_md.encode("utf-8"),
                    file_name=f"{docx_title}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

            with dl3:
                all_done = all(k in st.session_state for k in
                               ["questions","activities","lessonplan","eval_parent"])
                st.markdown('<div class="ppt-wrap">', unsafe_allow_html=True)
                if all_done:
                    if st.button("🎞️ PPT 생성 (Canva·발표용)", use_container_width=True, key="btn_ppt"):
                        with st.spinner("🖍️ PPT 슬라이드 만드는 중..."):
                            pptx_bytes = make_pptx(
                                grade, theme, book, lesson_time,
                                st.session_state.get("questions", {}),
                                st.session_state.get("activities", ""),
                                st.session_state.get("eval_parent", ""),
                            )
                        if pptx_bytes:
                            st.session_state["pptx_bytes"] = pptx_bytes
                else:
                    missing = sum(1 for k in ["questions","activities","lessonplan","eval_parent"]
                                  if k not in st.session_state)
                    st.button(
                        f"🎞️ PPT 생성 (수업안 {4-missing}/4 완료)",
                        use_container_width=True, key="btn_ppt",
                        disabled=True,
                        help="'수업안 전체 생성' 버튼을 먼저 눌러 주세요"
                    )
                if "pptx_bytes" in st.session_state:
                    st.download_button(
                        "⬇️ PPT 다운로드",
                        data=st.session_state["pptx_bytes"],
                        file_name=f"{book}_{theme}_수업PPT.pptx",
                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                        use_container_width=True,
                    )
                st.markdown("</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
